"""
ChromaDB Service for Document Embeddings.

This service handles:
1. Document chunking
2. Embedding generation
3. Storage in ChromaDB
4. Hybrid retrieval (Vector + BM25)
"""

import os
import logging
from typing import List, Dict, Any, Optional, Tuple
from pathlib import Path
import chromadb
from chromadb.config import Settings
from langchain_core.documents import Document
from langchain_text_splitters import RecursiveCharacterTextSplitter

# Configuration: Set to False to disable Hugging Face transformers
# This prevents SSL certificate issues and avoids downloading models from HuggingFace
USE_HUGGINGFACE_TRANSFORMERS = os.getenv('USE_HUGGINGFACE_TRANSFORMERS', 'true').lower() == 'true'

# Optional sentence_transformers - will fallback to OpenAI if not available
SENTENCE_TRANSFORMERS_AVAILABLE = False
if USE_HUGGINGFACE_TRANSFORMERS:
    try:
        from sentence_transformers import SentenceTransformer
        SENTENCE_TRANSFORMERS_AVAILABLE = True
    except ImportError:
        SENTENCE_TRANSFORMERS_AVAILABLE = False

from rank_bm25 import BM25Okapi
import numpy as np
from collections import defaultdict

from services.openai_service import get_openai_service
from services.document_loader import get_document_loader

logger = logging.getLogger(__name__)


class ChromaDBService:
    """Service for managing document embeddings in ChromaDB with hybrid search."""
    
    def __init__(
        self,
        collection_name: str = "finance_documents",
        persist_directory: str = "./chroma_db",
        embedding_model: str = "all-MiniLM-L6-v2",
        chunk_size: int = 1000,
        chunk_overlap: int = 200
    ):
        self.collection_name = collection_name
        self.persist_directory = persist_directory
        self.chunk_size = chunk_size
        self.chunk_overlap = chunk_overlap
        
        # Initialize ChromaDB client
        self.client = chromadb.PersistentClient(
            path=persist_directory,
            settings=Settings(anonymized_telemetry=False)
        )
        
        # Get or create collection
        self.collection = self.client.get_or_create_collection(
            name=collection_name,
            metadata={"hnsw:space": "cosine"}
        )
        
        # Initialize embedding model
        self.embedding_model = None
        self.openai_service = None
        
        if USE_HUGGINGFACE_TRANSFORMERS and SENTENCE_TRANSFORMERS_AVAILABLE:
            try:
                self.embedding_model = SentenceTransformer(embedding_model)
                logger.info(f"Initialized embedding model: {embedding_model}")
            except Exception as e:
                logger.warning(f"Failed to initialize SentenceTransformer: {e}")
                self.embedding_model = None
        else:
            logger.info("Hugging Face transformers disabled. Will use OpenAI embeddings if API key is available.")
        
        # Initialize OpenAI service as fallback or primary embedding source
        if not self.embedding_model:
            try:
                self.openai_service = get_openai_service()
                # Check if API key is available
                if not self.openai_service.api_key:
                    logger.warning("OPENAI_API_KEY not set. Set USE_HUGGINGFACE_TRANSFORMERS=true to use local embeddings or set OPENAI_API_KEY.")
                    self.openai_service = None
                else:
                    logger.info("Using OpenAI embeddings as embedding source")
            except Exception as e2:
                logger.warning(f"Failed to initialize OpenAI service: {e2}")
                self.openai_service = None
        
        # Initialize text splitter
        self.text_splitter = RecursiveCharacterTextSplitter(
            chunk_size=chunk_size,
            chunk_overlap=chunk_overlap,
            length_function=len,
        )
        
        # Initialize document loader
        self.document_loader = get_document_loader()
        
        # BM25 index (in-memory for hybrid search)
        self.bm25_index = None
        self.bm25_documents = []
        
        logger.info(f"ChromaDB service initialized - Collection: {collection_name}")
    
    def add_document(
        self,
        file_path: str,
        document_id: str,
        document_name: str,
        metadata: Optional[Dict[str, Any]] = None
    ) -> Dict[str, Any]:
        """
        Process a document and add it to ChromaDB.
        
        Args:
            file_path: Path to the document file
            document_id: Unique identifier for the document
            document_name: Name of the document
            metadata: Additional metadata
        
        Returns:
            Dictionary with processing results including chunks
        """
        try:
            # Load document
            logger.info(f"Loading document: {document_name}")
            documents = self._load_document(file_path)
            
            if not documents:
                raise ValueError(f"No content extracted from {file_path}")
            
            # Add document metadata
            for doc in documents:
                doc.metadata.update({
                    "document_id": str(document_id),
                    "document_name": document_name,
                    **(metadata or {})
                })
            
            # Split into chunks
            logger.info(f"Splitting document into chunks")
            chunks = self.text_splitter.split_documents(documents)
            
            # Generate embeddings and add to ChromaDB
            logger.info(f"Generating embeddings for {len(chunks)} chunks")
            chunk_data = self._prepare_chunks_for_chromadb(chunks, document_id)
            
            # Add to ChromaDB
            if chunk_data['ids']:
                self.collection.add(
                    ids=chunk_data['ids'],
                    embeddings=chunk_data['embeddings'],
                    documents=chunk_data['documents'],
                    metadatas=chunk_data['metadatas']
                )
                logger.info(f"Added {len(chunk_data['ids'])} chunks to ChromaDB")
            
            # Rebuild BM25 index with all documents
            self._rebuild_bm25_index()
            
            return {
                "document_id": str(document_id),
                "document_name": document_name,
                "chunks_count": len(chunks),
                "chunks": [
                    {
                        "chunk_id": chunk_id,
                        "content": content,
                        "metadata": meta,
                        "chunk_index": i
                    }
                    for i, (chunk_id, content, meta) in enumerate(
                        zip(chunk_data['ids'], chunk_data['documents'], chunk_data['metadatas'])
                    )
                ],
                "status": "success"
            }
            
        except Exception as e:
            logger.error(f"Error adding document to ChromaDB: {e}")
            raise
    
    def _load_document(self, file_path: str) -> List[Document]:
        """Load document using appropriate loader."""
        ext = Path(file_path).suffix.lower()
        
        if ext == '.pdf':
            return self.document_loader.load_pdf_file(file_path)
        elif ext == '.txt':
            return self.document_loader.load_text_file(file_path)
        elif ext in ['.docx', '.doc']:
            # For DOCX, read as text for now
            try:
                from docx import Document as DocxDocument
                doc = DocxDocument(file_path)
                text = '\n'.join([para.text for para in doc.paragraphs])
                return [Document(page_content=text, metadata={"source": file_path})]
            except Exception as e:
                logger.warning(f"Failed to load DOCX, trying as text: {e}")
                return self.document_loader.load_text_file(file_path)
        else:
            return self.document_loader.load_text_file(file_path)
    
    def _prepare_chunks_for_chromadb(
        self,
        chunks: List[Document],
        document_id: str
    ) -> Dict[str, List]:
        """Prepare chunks for ChromaDB storage."""
        ids = []
        embeddings = []
        documents = []
        metadatas = []
        
        for i, chunk in enumerate(chunks):
            chunk_id = f"{document_id}_chunk_{i}"
            ids.append(chunk_id)
            documents.append(chunk.page_content)
            
            # Generate embedding
            embedding = self._generate_embedding(chunk.page_content)
            embeddings.append(embedding)
            
            # Prepare metadata
            metadata = {
                **chunk.metadata,
                "chunk_index": i,
                "chunk_id": chunk_id
            }
            metadatas.append(metadata)
        
        return {
            "ids": ids,
            "embeddings": embeddings,
            "documents": documents,
            "metadatas": metadatas
        }
    
    def _generate_embedding(self, text: str) -> List[float]:
        """Generate embedding for text."""
        if self.embedding_model:
            return self.embedding_model.encode([text])[0].tolist()
        elif self.openai_service and self.openai_service.api_key:
            # Use OpenAI embeddings as fallback
            try:
                embeddings = self.openai_service.get_langchain_embeddings()
                return embeddings.embed_query(text)
            except Exception as e:
                logger.error(f"OpenAI embedding generation failed: {e}")
                raise ValueError(f"OpenAI embedding generation failed: {e}")
        else:
            raise ValueError(
                "No embedding model available. Either install sentence-transformers "
                "(pip install sentence-transformers) or set OPENAI_API_KEY environment variable."
            )
    
    def _rebuild_bm25_index(self):
        """Rebuild BM25 index from all documents in ChromaDB."""
        try:
            # Get all documents from ChromaDB
            results = self.collection.get(include=['documents', 'metadatas'])
            
            if results and results['documents']:
                # Create Document objects
                self.bm25_documents = [
                    Document(page_content=doc, metadata=meta)
                    for doc, meta in zip(results['documents'], results['metadatas'])
                ]
                
                # Build BM25 index
                tokenized_docs = [self._tokenize(doc.page_content) for doc in self.bm25_documents]
                self.bm25_index = BM25Okapi(tokenized_docs)
                
                logger.info(f"Rebuilt BM25 index with {len(self.bm25_documents)} documents")
        except Exception as e:
            logger.error(f"Failed to rebuild BM25 index: {e}")
    
    def hybrid_search(
        self,
        query: str,
        k: int = 5,
        vector_weight: float = 0.7,
        bm25_weight: float = 0.3,
        filter_metadata: Optional[Dict[str, Any]] = None
    ) -> List[Dict[str, Any]]:
        """
        Perform hybrid search combining vector similarity and BM25.
        
        Args:
            query: Search query
            k: Number of results to return
            vector_weight: Weight for vector search results (0-1)
            bm25_weight: Weight for BM25 search results (0-1)
            filter_metadata: Optional metadata filters
        
        Returns:
            List of search results with scores
        """
        try:
            # Vector search
            vector_results = self._vector_search(query, k * 2, filter_metadata)
            
            # BM25 search
            bm25_results = self._bm25_search(query, k * 2, filter_metadata)
            
            # Combine results
            hybrid_results = self._combine_results(
                vector_results, bm25_results, vector_weight, bm25_weight, k
            )
            
            logger.info(f"Hybrid search: {len(hybrid_results)} results for query '{query}'")
            return hybrid_results
            
        except Exception as e:
            logger.error(f"Hybrid search failed: {e}")
            # Fallback to vector search only
            return self._vector_search(query, k, filter_metadata)
    
    def _vector_search(
        self,
        query: str,
        k: int,
        filter_metadata: Optional[Dict[str, Any]] = None
    ) -> List[Dict[str, Any]]:
        """Perform vector similarity search."""
        try:
            query_embedding = self._generate_embedding(query)
            
            results = self.collection.query(
                query_embeddings=[query_embedding],
                n_results=k,
                where=filter_metadata,
                include=['documents', 'metadatas', 'distances']
            )
            
            search_results = []
            if results and results['ids']:
                for i, (doc_id, doc, meta, distance) in enumerate(zip(
                    results['ids'][0],
                    results['documents'][0],
                    results['metadatas'][0],
                    results['distances'][0]
                )):
                    search_results.append({
                        "id": doc_id,
                        "content": doc,
                        "metadata": meta,
                        "score": 1 - distance,  # Convert distance to similarity
                        "rank": i + 1,
                        "search_type": "vector"
                    })
            
            return search_results
            
        except Exception as e:
            logger.error(f"CromaDB service, Vector search failed: {e}")
            return []
    
    def _bm25_search(
        self,
        query: str,
        k: int,
        filter_metadata: Optional[Dict[str, Any]] = None
    ) -> List[Dict[str, Any]]:
        """Perform BM25 keyword search."""
        try:
            if not self.bm25_index or not self.bm25_documents:
                logger.warning("BM25 index not available")
                return []
            
            # Tokenize query
            query_tokens = self._tokenize(query)
            
            # Get BM25 scores
            bm25_scores = self.bm25_index.get_scores(query_tokens)
            
            # Sort by score
            sorted_indices = np.argsort(bm25_scores)[::-1]
            
            # Filter by metadata if provided
            search_results = []
            for idx in sorted_indices:
                if len(search_results) >= k:
                    break
                
                doc = self.bm25_documents[idx]
                
                # Apply metadata filter
                if filter_metadata:
                    matches = all(
                        doc.metadata.get(key) == value
                        for key, value in filter_metadata.items()
                    )
                    if not matches:
                        continue
                
                search_results.append({
                    "id": doc.metadata.get("chunk_id", f"chunk_{idx}"),
                    "content": doc.page_content,
                    "metadata": doc.metadata,
                    "score": float(bm25_scores[idx]),
                    "rank": len(search_results) + 1,
                    "search_type": "bm25"
                })
            
            return search_results
            
        except Exception as e:
            logger.error(f"BM25 search failed: {e}")
            return []
    
    def _combine_results(
        self,
        vector_results: List[Dict[str, Any]],
        bm25_results: List[Dict[str, Any]],
        vector_weight: float,
        bm25_weight: float,
        k: int
    ) -> List[Dict[str, Any]]:
        """Combine vector and BM25 results with weighted scoring."""
        # Create combined score mapping
        combined_scores = defaultdict(lambda: {"vector": 0.0, "bm25": 0.0, "result": None})
        
        # Normalize and add vector scores
        if vector_results:
            max_vector = max(r["score"] for r in vector_results)
            for result in vector_results:
                doc_id = result["id"]
                normalized_score = result["score"] / max_vector if max_vector > 0 else 0
                combined_scores[doc_id]["vector"] = normalized_score
                combined_scores[doc_id]["result"] = result
        
        # Normalize and add BM25 scores
        if bm25_results:
            max_bm25 = max(r["score"] for r in bm25_results)
            for result in bm25_results:
                doc_id = result["id"]
                normalized_score = result["score"] / max_bm25 if max_bm25 > 0 else 0
                combined_scores[doc_id]["bm25"] = normalized_score
                if combined_scores[doc_id]["result"] is None:
                    combined_scores[doc_id]["result"] = result
        
        # Calculate hybrid scores
        hybrid_results = []
        for doc_id, scores in combined_scores.items():
            if scores["result"]:
                hybrid_score = (
                    vector_weight * scores["vector"] +
                    bm25_weight * scores["bm25"]
                )
                
                result = scores["result"].copy()
                result["hybrid_score"] = hybrid_score
                result["vector_score"] = scores["vector"]
                result["bm25_score"] = scores["bm25"]
                result["search_type"] = "hybrid"
                hybrid_results.append(result)
        
        # Sort by hybrid score and return top k
        hybrid_results.sort(key=lambda x: x["hybrid_score"], reverse=True)
        return hybrid_results[:k]
    
    def _tokenize(self, text: str) -> List[str]:
        """Simple tokenization for BM25."""
        return text.lower().split()
    
    def get_document_chunks(self, document_id: str) -> List[Dict[str, Any]]:
        """Get all chunks for a specific document."""
        try:
            results = self.collection.get(
                where={"document_id": str(document_id)},
                include=['documents', 'metadatas']
            )
            
            chunks = []
            if results and results['ids']:
                for chunk_id, content, metadata in zip(
                    results['ids'],
                    results['documents'],
                    results['metadatas']
                ):
                    chunks.append({
                        "chunk_id": chunk_id,
                        "content": content,
                        "metadata": metadata,
                        "chunk_index": metadata.get("chunk_index", 0)
                    })
                
                # Sort by chunk index
                chunks.sort(key=lambda x: x["chunk_index"])
            
            return chunks
            
        except Exception as e:
            logger.error(f"Error getting document chunks: {e}")
            return []
    
    def delete_document(self, document_id: str) -> bool:
        """Delete all chunks for a document from ChromaDB."""
        try:
            # Get all chunk IDs for this document
            results = self.collection.get(
                where={"document_id": str(document_id)},
                include=['documents']
            )
            
            if results and results['ids']:
                self.collection.delete(ids=results['ids'])
                logger.info(f"Deleted {len(results['ids'])} chunks for document {document_id}")
                
                # Rebuild BM25 index
                self._rebuild_bm25_index()
                
                return True
            
            return False
            
        except Exception as e:
            logger.error(f"Error deleting document from ChromaDB: {e}")
            return False
    
    def get_collection_stats(self) -> Dict[str, Any]:
        """Get statistics about the ChromaDB collection."""
        try:
            count = self.collection.count()
            
            return {
                "collection_name": self.collection_name,
                "total_chunks": count,
                "embedding_model": str(self.embedding_model) if self.embedding_model else "OpenAI",
                "chunk_size": self.chunk_size,
                "chunk_overlap": self.chunk_overlap,
                "bm25_enabled": self.bm25_index is not None
            }
        except Exception as e:
            logger.error(f"Error getting collection stats: {e}")
            return {}


# Singleton instance
_chromadb_service = None


def get_chromadb_service(collection_name: str = "finance_documents") -> ChromaDBService:
    """Get or create ChromaDB service instance."""
    global _chromadb_service
    
    if _chromadb_service is None or _chromadb_service.collection_name != collection_name:
        _chromadb_service = ChromaDBService(collection_name=collection_name)
    
    return _chromadb_service

